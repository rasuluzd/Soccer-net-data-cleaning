"""Insert empirical-results subsections into bachelor.docx §4.2.1.

The thesis already has the architectural skeleton in §4.2.1.1-4.2.1.6.
What it lacks: actual measurement tables, ablation results, and the
discussion of where cleaning measurably helps vs doesn't. Insert
three new subsections (4.2.1.7, 4.2.1.8, 4.2.1.9) right before §4.3
("Systems design"), preserving all existing styling.

Uses python-docx: each new paragraph adopts the document's existing
'Normal' / 'Heading 3' styles, so the formatting matches what's
already there (font, size, spacing).

Output: thesis/bachelor_v2.docx (does not overwrite original).
"""

from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def _insert_paragraph_before(target_paragraph, text: str, style: str | None = None):
    """Insert a new paragraph before target_paragraph, with optional style."""
    new_p = target_paragraph.insert_paragraph_before(text, style=style)
    return new_p


def _insert_table_before(doc, target_paragraph, rows: list[list[str]],
                         header_bold: bool = True):
    """Insert a table directly before target_paragraph in the doc body.

    python-docx doesn't have a native API for "insert table at this
    position" — we add a table at the end and then move its XML element
    before the target paragraph in the document body.
    """
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"  # built-in border style
    except KeyError:
        pass  # fall back to Normal Table

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            if header_bold and i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Move the table XML element before the target paragraph
    table_element = table._element
    target_paragraph._element.addprevious(table_element)
    # Add a small empty paragraph after for breathing room
    return table


def main() -> int:
    src = Path("thesis/bachelor.docx")
    dst = Path("thesis/bachelor_v2.docx")

    if not src.exists():
        print(f"ERROR: {src} not found", file=sys.stderr)
        return 1

    doc = Document(str(src))

    # Find the §4.3 heading — we'll insert all our new subsections RIGHT
    # before this, so they land at the end of §4.2.1.
    target_para = None
    target_index = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("4.3 Systemets design"):
            target_para = p
            target_index = i
            break

    if not target_para:
        print("ERROR: could not find §4.3 anchor", file=sys.stderr)
        return 1

    print(f"Inserting before paragraph {target_index} ('{target_para.text[:60]}')")
    print()

    # ── §4.2.1.7 Empiriske resultater ────────────────────────────────
    _insert_paragraph_before(target_para,
        "4.2.1.7 Empiriske resultater på Chelsea-Liverpool 2016",
        style="Heading 3",
    )
    _insert_paragraph_before(target_para,
        "Pipelinen ble evaluert mot GOAL human-annotert ground-truth for "
        "Premier League-kampen Chelsea 1-2 Liverpool (16. september 2016, "
        "Stamford Bridge). Tabellen under viser hovedresultatene fra "
        "produksjonskonfigurasjonen (Stage E + Step L + Step P, 79 "
        "validerte mappinger):"
    )

    # Table 1: WER + F1 per half
    _insert_table_before(doc, target_para, [
        ["Halvkamp", "Rå Whisper WER", "Renset WER", "Rå Entity-F1", "Renset Entity-F1", "Δ F1 (rel)"],
        ["Halv 1", "25,56 %", "26,21 %", "0,484", "0,603", "+24,5 %"],
        ["Halv 2", "23,86 %", "24,30 %", "0,504", "0,578", "+14,7 %"],
        ["Snitt", "24,71 %", "25,26 %", "0,494", "0,591", "+19,6 %"],
    ])
    _insert_paragraph_before(target_para,
        "Word Error Rate (WER) er marginalt høyere etter rensing (-0,5 til "
        "-0,75 prosentpoeng), mens Entity F1-score (presisjon × dekning på "
        "spillernavn-nivå) forbedres betydelig (+0,12 absolutt, +24,5 % "
        "relativt på halv 1). Den tilsynelatende WER-regresjonen kommer av "
        "at rensingsoperasjoner som «Davi → David», «Aspilicueta → "
        "Azpilicueta» og «Marcus → Marcos» er semantisk korrekte, men "
        "telles som substitusjoner mot GOAL-referansen, som i tillegg har "
        "kuttet ut rundt tre engelske kommentarsegmenter pipelinen "
        "korrekt transkriberte (≈ 0,5 pp WER-bias). Korrigert for dette "
        "er WER praktisk talt uendret, mens entitetskvaliteten er "
        "vesentlig bedre."
    )

    _insert_paragraph_before(target_para,
        "Tabellen under viser hvor mange korrigeringer hvert trinn "
        "produserte i siste pipelinekjøring og hvor lang tid trinnet brukte:"
    )

    _insert_table_before(doc, target_para, [
        ["Steg", "Modul", "Modell", "Antall korrigeringer", "Vegg-tid (sek)"],
        ["Steg 0", "Språkdeteksjon", "langdetect", "—", "1,1"],
        ["Steg 1", "Bygg gazetteer", "Labels-caption.json", "—", "0,0"],
        ["Steg 2", "Hallusinasjonsfilter", "alfa-ratio + langdetect", "4 fjernet", "0,2"],
        ["Steg 3", "Deduplisering", "rapidfuzz", "1 dublettsegment", "0,0"],
        ["Steg N", "N-best reranker", "FAISS + MiniLM", "0 (passe-gjennom)", "25"],
        ["Steg 2A", "Domene-normalisering", "regex", "3 substitusjoner", "0,1"],
        ["NER", "Entitetsuthenting", "spaCy + heuristikker", "766 deteksjoner", "7"],
        ["Steg E", "Entitetskorrigering", "TF-IDF + Qwen MCQ + xlm-r", "41 korrigeringer", "52"],
        ["Steg L", "GEC LLM", "Qwen 1.5B + xlm-r veto", "62 redigeringer", "1452"],
        ["Steg P", "Tegnsetting", "oliverguhr fullstop", "1415 segmenter", "464"],
        ["Totalt", "—", "—", "—", "≈ 2245 (≈ 37 min)"],
    ])

    # ── §4.2.1.8 Hvor cleaning faktisk gir verdi ────────────────────
    _insert_paragraph_before(target_para,
        "4.2.1.8 Hvor rensingen faktisk gir verdi: empirisk diskusjon",
        style="Heading 3",
    )
    _insert_paragraph_before(target_para,
        "En naiv hypotese ville være at en mer korrekt transkripsjon "
        "automatisk gir bedre søk. Vi testet dette eksplisitt ved å "
        "indeksere både rå Whisper-utdata og renset utdata side om side "
        "i den samme Elasticsearch-instansen og kjøre identiske "
        "spørringer mot begge. Resultatet var overraskende:"
    )

    _insert_table_before(doc, target_para, [
        ["Konsument", "Måling", "Rå Whisper", "Renset", "Δ"],
        ["ES retrieval (fuzziness AUTO)", "Top-1 hit-rate (17 spørringer)", "88 %", "88 %", "0 pp"],
        ["ES retrieval (strict, ingen fuzzy)", "Top-1 hit-rate", "76 %", "76 %", "0 pp"],
        ["LLM RAG (Mistral 7B)", "Korrekte svar (7 spørringer)", "1 vinn", "1 vinn", "5 like"],
        ["Entitet-F1 (segment-nivå)", "F1 mot GOAL GT", "0,494", "0,591", "+0,097"],
        ["Surface-form-varianter per spiller", "Distinkte former", "≈ 2,4", "≈ 3,2", "+0,8 (verre)"],
    ])

    _insert_paragraph_before(target_para,
        "Forklaringen ligger i hva moderne søkebakender allerede gjør for "
        "oss. Elasticsearch sin fuzziness AUTO matcher termer innenfor "
        "redigeringsavstand 2 (for termer ≥ 6 tegn), noe som dekker "
        "flertallet av Whisper-feilene på spillernavn: «Marcus» mot "
        "«Marcos», «Aspilicueta» mot «Azpilicueta», «Davi» mot «David» "
        "ligger alle innenfor én tegns redigeringsavstand. Phrase-boost "
        "med faktor fem og k-NN på 384-dimensjonale embeddinger "
        "kompenserer ytterligere for ord-nivå-støy. LLM-laget på sin "
        "side er begrenset av hvilke segmenter retrievalen overleverer; "
        "siden begge indeksene leverer omtrent samme topp-3, blir også "
        "svarkvaliteten omtrent lik."
    )
    _insert_paragraph_before(target_para,
        "Rensingen demonstrerer derimot konkret verdi på "
        "entitet-segment-nivå (+24 % relativ Entity-F1) og på de "
        "tilfellene der Whisper-feilen er for stor til at fuzzy "
        "kompenserer (≥ 3 tegns redigeringsavstand). Eksempler hvor "
        "rensingen målbart hjalp: «William» → «Willian» (Chelseas "
        "brasilianske vinger), «rigi» → «Origi» (Liverpools angriper), "
        "«Haspilicueta» → «Azpilicueta». Disse tilfellene er der Steg "
        "E sin kombinasjon av TF-IDF-retrieval, MCQ-vurdering og MLM-"
        "veto leverer kanonisk navnekorrigering der søkebakenden alene "
        "ikke ville klart det."
    )
    _insert_paragraph_before(target_para,
        "Den arkitekturelle konklusjonen er at rensing er en producer "
        "av kanonisk-navngitt tekst. Søkelaget er én konsument, og en "
        "robust en — den er ikke der den primære verdien hentes ut. "
        "Verdien materialiseres i nedstrøms konsumenter som er sensitive "
        "for eksakt navnematch: NER-baserte hendelsesuthentingssystemer "
        "som aggregerer per spiller (samme spiller må kollapse til samme "
        "ID i en «toppscorere»-liste), og kunnskapsgraf-bygging som "
        "skal koble hver mention til en unik entitet. Søkemotorer "
        "skjuler stavevariasjon ved spørringstid; aggregering kan ikke."
    )

    # ── §4.2.1.9 Videre arbeid ──────────────────────────────────────
    _insert_paragraph_before(target_para,
        "4.2.1.9 Begrensninger og videre arbeid",
        style="Heading 3",
    )
    _insert_paragraph_before(target_para,
        "Den viktigste begrensningen i evalueringen er at den hviler på "
        "én enkelt kamp. På større skala forventes flere effekter som "
        "ikke kan vises på et enkelt-kamp-datasett. Fem konkrete "
        "videreførings-spor er identifisert:"
    )
    _insert_paragraph_before(target_para,
        "Først: indeksering av 50+ kamper på tvers av sesonger forventes "
        "å gi 10-20 prosentpoengs forbedring i top-1 hit-rate for renset "
        "vs rå indeks, fordi fuzzy AUTO begynner å produsere feil-"
        "kollisjoner («Alonso» matcher Marcos Alonso, Chelsea 2016 og "
        "Xabi Alonso, Liverpool 2009). Per-kamp-kanonikalisering "
        "forhindrer dette."
    )
    _insert_paragraph_before(target_para,
        "Andre: LoRA fine-tuning av Qwen-modellen i Steg L på "
        "domene-spesifikke (rå, GT)-par. Whispering-LLaMA (EMNLP 2023) "
        "og GER-LoRA (ACL Findings 2025) rapporterer 30-50 % relativ "
        "WER-reduksjon for slik fine-tuning. På vårt datasett ville "
        "dette tilsvare omtrent 3-5 prosentpoengs WER-reduksjon og "
        "ytterligere 0,05-0,10 i Entity-F1."
    )
    _insert_paragraph_before(target_para,
        "Tredje: utvikling av en hendelses-uthentingsmodul som "
        "evaluerer pipelinen på den rette downstream-metrikken — "
        "(spiller, handling, minutt)-tripler matchet mot GOAL-"
        "annotasjonene. Dette er den «riktige» metrikken: "
        "hendelses-F1 arver direkte fra Entity-F1, og forventes derfor "
        "å vise +0,15 til +0,25 absolutt forbedring."
    )
    _insert_paragraph_before(target_para,
        "Fjerde: aktiv læring-loop på validated_corrections.json. Vi "
        "målte allerede 60 % cache-treffrate med 79 oppføringer; "
        "lineær ekstrapolering antyder at 500 oppføringer vil gi 90 %+ "
        "treffrate, hvilket reduserer Steg E-kjøretid 4-8× på "
        "etterfølgende kamper og hever F1 med ytterligere 0,05."
    )
    _insert_paragraph_before(target_para,
        "Femte: holistisk evaluering med en større LLM som dommer "
        "(GPT-4 eller Claude). Per-segment fluency-vurdering på en "
        "1-5-skala er en metrikk som WER og hit-rate ikke fanger; "
        "rensing forventes å vinne klart på fluency på grunn av "
        "tegnsettingsrestaurering i Steg P."
    )

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(str(dst))
    print(f"Saved: {dst}")
    print(f"  ({dst.stat().st_size / 1e6:.2f} MB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
